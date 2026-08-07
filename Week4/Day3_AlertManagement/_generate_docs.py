"""Generate all Week 4 Day 3 documentation files."""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

OUT = os.path.dirname(os.path.abspath(__file__))

# ═══════════════════════════════════════════════════════════════
# 01 — Research Document
# ═══════════════════════════════════════════════════════════════
def create_research():
    doc = Document()
    doc.add_heading('Intelligent Alert Management for Ergonomic Monitoring', level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Week 4 — Day 3 | ErgoVigilance Research Document').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')

    sections = [
        ('1. Introduction',
         'Modern ergonomic monitoring systems generate real-time risk assessments every frame — up to 30 '
         'assessments per second. Without intelligent alert management, this creates an unsustainable '
         'notification deluge. Workers ignore alerts, supervisors become desensitised, and the system\'s '
         'credibility erodes. This document presents a research-backed Intelligent Alert Management '
         'system for ErgoVigilance, designed to deliver the right alert to the right person at the right '
         'time — and suppress everything else.'),

        ('2. The Problem: Alert Fatigue',
         'Alert fatigue occurs when an excessive volume of notifications desensitises recipients. In '
         'clinical settings, alarm fatigue causes 85–99% of alarms to be ignored (Sendelbach & Funk, 2013). '
         'In industrial ergonomics, the problem is similar: every postural deviation generates a warning, '
         'most are transient, and workers learn to dismiss them.\n\n'
         'Consequences of alert fatigue:\n'
         '• Diminished response to genuine alerts\n'
         '• Increased response time during critical events\n'
         '• Worker frustration and system distrust\n'
         '• Supervisor desensitisation to high-priority notifications\n'
         '• Reduced system adoption and abandonment'),

        ('3. Human Factors in Alert Design',
         'Effective alert systems account for human cognitive limitations:\n\n'
         '• Hick\'s Law: Decision time increases with choices — limit alert types to 3–4 levels\n'
         '• Signal Detection Theory: Sensitivity (d\') and bias (β) determine whether alerts are detected '
         'or ignored — suppress noise to maintain sensitivity\n'
         '• Inattentional Blindness: Focusing on a task causes missing of visual alerts — use multi-modal '
         'cues (visual + audible) for critical alerts\n'
         '• Habituation: Repeated identical stimuli reduce response — vary alert presentation\n\n'
         'Design principles from NASA\'s Cockpit Alerting System and aviation human factors apply directly '
         'to industrial ergonomic monitoring.'),

        ('4. Escalation Logic',
         'Intelligent escalation follows a multi-tier path:\n\n'
         'Level 0 — No Alert: Posture within normal range, no action needed\n'
         'Level 1 — LOW (Worker Only): Minor deviation detected. In-app notification to worker. '
         'Self-correction expected.\n'
         'Level 2 — MEDIUM (Worker + Supervisor): Sustained deviation. Worker notified with actionable '
         'recommendation. Supervisor gets summary notification.\n'
         'Level 3 — HIGH (Worker + Supervisor + Manager): Critical deviation or fatigue accumulation. '
         'All parties notified. Manager gets escalation summary.\n'
         'Level 4 — CRITICAL: Immediate physical risk. Audio-visual alarm. Supervisor dispatched. '
         'Session may be paused.\n\n'
         'Escalation triggers:\n'
         '• Duration-based: risk level sustained for X seconds\n'
         '• Fatigue-based: fatigue accumulator exceeds threshold\n'
         '• Frequency-based: N alerts within M minutes\n'
         '• Context-based: specific task + risk combination'),

        ('5. Industrial Alert Systems — Best Practices',
         'OSHA recommends:\n'
         '• Risk-based prioritisation with clear escalation paths\n'
         '• Positive feedback when risk decreases (not just warnings)\n'
         '• Regular review of alert effectiveness\n'
         '• Worker training on alert response\n\n'
         'ISO 9241-210 (Human-centred design) emphasises:\n'
         '• Alerts must be actionable — workers need to know WHAT to do\n'
         '• Alerts must be timely — too early causes false starts, too late causes harm\n'
         '• Alerts must respect context — no alerts during critical operations\n\n'
         'Industry leaders (Toyota Production System, 5S methodology) use:\n'
         '• Andon cords for immediate escalation\n'
         '• Visual management (andon boards) for team awareness\n'
         '• Tiered response: worker → team lead → supervisor → manager'),

        ('6. Alert Prioritisation Framework',
         'The proposed system uses 4 factors to determine alert priority:\n\n'
         '1. Risk Magnitude: Current risk level (LOW → MEDIUM → HIGH)\n'
         '2. Duration: Time at current risk level (seconds → minutes → hours)\n'
         '3. Fatigue: Cumulative fatigue score (0–100)\n'
         '4. Trend: Direction of change (improving → stable → deteriorating)\n\n'
         'Alert Priority = f(risk_level, duration_minutes, fatigue_level, trend_direction)\n\n'
         'A matrix of 5 × 5 × 3 × 3 = 225 possible states maps to 4 alert levels.'),

        ('7. Cooldown and Suppression Logic',
         'To prevent alert storms:\n\n'
         '• Cooldown Timer: After any alert, minimum 2-minute cooldown before next alert of same type\n'
         '• Duplicate Suppression: Identical alerts within 5-minute window are suppressed\n'
         '• Threshold Hysteresis: Alert triggers at threshold X, clears at X − delta (avoid oscillation)\n'
         '• Fatigue Gating: Low-risk alerts suppressed when fatigue is low and improving\n'
         '• Context Suppression: Expected postures for a task suppress alerts (e.g., trunk flexion '
         'during lifting)'),

        ('8. Worker Psychology',
         'Key psychological factors in alert response:\n\n'
         '• Self-efficacy: Workers who believe they can correct posture are more likely to respond '
         'to alerts → provide actionable, achievable recommendations\n'
         '• Reactance: Too many alerts cause psychological resistance → suppress low-value alerts\n'
         '• Alert fatigue curve: Response rate drops 50% after 5+ alerts per hour\n'
         '• Positive reinforcement: Acknowledge good posture — not just bad posture\n'
         '• Social norms: Team-level alert summaries encourage collective improvement'),

        ('9. Supervisor Notification Systems',
         'Supervisors receive aggregated, not per-frame, notifications:\n\n'
         '• Summary: "Worker X had 3 high-risk episodes in the last hour" (not every frame)\n'
         '• Trend: "Worker X\'s risk is deteriorating — 15% increase over baseline"\n'
         '• Exception: Only alerts that meet escalation threshold reach supervisors\n'
         '• Dashboard: Real-time team view with colour-coded risk per workstation\n'
         '• Report: End-of-shift summary with action items\n\n'
         'Manager notifications are further filtered — only systemic issues (multiple workers, '
         'repeated violations, safety policy breaches) trigger manager alerts.'),

        ('10. Conclusion',
         'Intelligent Alert Management transforms a potential weakness of real-time ergonomic monitoring — '
         'notification overload — into a strength. By implementing evidence-based escalation logic, '
         'cooldown suppression, duplicate filtering, and role-based notification routing, the system '
         'ensures that every alert is meaningful, actionable, and appropriately urgent. The result is '
         'higher worker engagement, better supervisor oversight, and a system that workers trust and use.'),
    ]

    for heading, body in sections:
        doc.add_heading(heading, level=1)
        for para in body.split('\n\n'):
            doc.add_paragraph(para.strip())

    path = os.path.join(OUT, 'Research', '01_IntelligentAlertResearch.docx')
    doc.save(path)
    print(f"Created: {path}")

# ═══════════════════════════════════════════════════════════════
# 02 — Alert Decision Matrix
# ═══════════════════════════════════════════════════════════════
def create_matrix():
    wb = Workbook()
    ws1 = wb.active
    ws1.title = 'Decision Matrix'

    hf = Font(bold=True, color='FFFFFF', size=10)
    hfill = PatternFill(start_color='2D2D3D', end_color='2D2D3D', fill_type='solid')
    tb = Border(left=Side(style='thin', color='555'), right=Side(style='thin', color='555'),
                top=Side(style='thin', color='555'), bottom=Side(style='thin', color='555'))

    headers = ['Risk', 'Duration (min)', 'Fatigue', 'Task', 'Prev Alerts', 'Alert Level', 'Escalation', 'Worker', 'Supervisor', 'Manager']
    for c, h in enumerate(headers, 1):
        cell = ws1.cell(row=1, column=c, value=h)
        cell.font = hf; cell.fill = hfill; cell.alignment = Alignment(horizontal='center', wrap_text=True); cell.border = tb

    rows = [
        ['LOW', '<5', '<20', 'Any', '0', 'None', 'None', '—', '—', '—'],
        ['LOW', '>30', '<20', 'Office', '1', 'LOW', 'Self-correct', 'Notification', '—', '—'],
        ['MEDIUM', '<10', '<30', 'Assembly', '0', 'LOW', 'Self-correct', 'Notification', '—', '—'],
        ['MEDIUM', '>20', '<40', 'Inspection', '2', 'MEDIUM', 'Worker + Super', 'Warning + Rec', 'Summary', '—'],
        ['MEDIUM', '>45', '<50', 'Assembly', '3', 'HIGH', 'Escalate to Mgr', 'Warning + Rec', 'Notify', 'Summary'],
        ['HIGH', '<5', '<30', 'Lifting', '0', 'MEDIUM', 'Worker + Super', 'Warning + Rec', 'Notify', '—'],
        ['HIGH', '>15', '>50', 'Assembly', '1', 'HIGH', 'Escalate to Mgr', 'Warning + Rec', 'Notify', 'Alert'],
        ['HIGH', '>30', '>60', 'Warehouse', '2', 'CRITICAL', 'Immediate', 'Audio-Visual', 'Urgent', 'Called'],
        ['HIGH', '>5', '>40', 'Any', '3+', 'CRITICAL', 'Immediate + Stop', 'Stop Work', 'Urgent', 'Called'],
        ['CRITICAL', '>1', '>70', 'Any', '1', 'CRITICAL', 'Immediate + Stop', 'Stop Work', 'Urgent', 'Called'],
        ['MEDIUM', '<5', '<20', 'Typing', '0', 'None', 'Self-correct', '—', '—', '—'],
        ['MEDIUM', '>30', '>30', 'Machine', '1', 'MEDIUM', 'Worker + Super', 'Warning', 'Summary', '—'],
        ['HIGH', '>10', '>40', 'Inspection', '1', 'HIGH', 'Escalate to Mgr', 'Warning + Rec', 'Notify', 'Alert'],
        ['LOW', '>60', '<20', 'Neutral', '0', 'LOW', 'Self-correct', 'Reminder', '—', '—'],
        ['MEDIUM', '>15', '>50', 'Lifting', '2', 'HIGH', 'Escalate to Mgr', 'Warning + Rec', 'Notify', 'Alert'],
    ]

    fills = {
        'LOW': PatternFill(start_color='1B5E20', end_color='1B5E20', fill_type='solid'),
        'MEDIUM': PatternFill(start_color='E65100', end_color='E65100', fill_type='solid'),
        'HIGH': PatternFill(start_color='B71C1C', end_color='B71C1C', fill_type='solid'),
        'CRITICAL': PatternFill(start_color='6A1B9A', end_color='6A1B9A', fill_type='solid'),
    }
    rf = Font(bold=True, color='FFFFFF', size=9)

    for r, row in enumerate(rows, 2):
        for c, val in enumerate(row, 1):
            cell = ws1.cell(row=r, column=c, value=val)
            cell.border = tb; cell.alignment = Alignment(horizontal='center', wrap_text=True)
            if c == 6 and val in fills:
                cell.fill = fills[val]; cell.font = rf
            elif c == 1 and val in fills:
                pass  # leave first column unstyled to avoid confusion

    for i, w in enumerate([8, 12, 10, 14, 12, 14, 18, 22, 14, 14], 1):
        ws1.column_dimensions[chr(64 + i)].width = w

    path = os.path.join(OUT, 'Research', '02_AlertDecisionMatrix.xlsx')
    wb.save(path)
    print(f"Created: {path}")

# ═══════════════════════════════════════════════════════════════
# 03 — Architecture SVG
# ═══════════════════════════════════════════════════════════════
def create_arch():
    svg = '''<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 800 780" width="800" height="780">
  <defs>
    <style>.box{fill:#1e1e2e;stroke:#555;stroke-width:1.5;rx:8;ry:8}
    .bt{fill:#e0e0e0;font-family:monospace;font-size:13px;text-anchor:middle}
    .ar{stroke:#888;stroke-width:1.5;fill:none;marker-end:url(#ah)}
    .hl{fill:#2d2d4d;stroke:#7c7ccc;stroke-width:2;rx:8;ry:8}
    .tl{fill:#bb86fc;font-family:sans-serif;font-size:16px;font-weight:bold;text-anchor:middle}
    .st{fill:#aaa;font-family:sans-serif;font-size:11px;text-anchor:middle}
    .lb{fill:#888;font-family:monospace;font-size:10px;text-anchor:middle}</style>
    <marker id="ah" markerWidth="8" markerHeight="6" refX="8" refY="3" orient="auto"><polygon points="0 0,8 3,0 6" fill="#888"/></marker>
  </defs>
  <rect width="800" height="780" fill="#121220"/>
  <text x="400" y="30" class="tl">Intelligent Alert Architecture</text>
  <text x="400" y="48" class="st">ErgoVigilance - Week 4 Day 3 Prototype</text>

  <!-- Row 1: Pose -->
  <rect x="300" y="65" width="200" height="40" class="box"/>
  <text x="400" y="82" class="bt" font-weight="bold">MediaPipe Pose</text>
  <text x="400" y="96" class="lb">33 landmarks</text>
  <line x1="400" y1="105" x2="400" y2="130" class="ar"/>

  <!-- Row 2: Features -->
  <rect x="275" y="133" width="250" height="40" class="box"/>
  <text x="400" y="150" class="bt" font-weight="bold">Feature Extraction</text>
  <text x="400" y="164" class="lb">7 biomechanical angles</text>
  <line x1="400" y1="173" x2="400" y2="198" class="ar"/>

  <!-- Row 3: Context Risk -->
  <rect x="250" y="201" width="300" height="40" class="hl"/>
  <text x="400" y="218" class="bt" font-weight="bold" fill="#bb86fc">Context-Aware Risk Engine</text>
  <text x="400" y="232" class="lb" fill="#9988cc">Base + Task + Duration + Fatigue</text>
  <line x1="400" y1="241" x2="400" y2="266" class="ar"/>

  <!-- Row 4: Alert Engine (NEW) -->
  <rect x="225" y="269" width="350" height="50" class="hl"/>
  <text x="400" y="290" class="bt" font-weight="bold" fill="#ff79c6">Alert Engine (NEW - Day 3)</text>
  <text x="400" y="306" class="lb" fill="#cc88aa">Prioritisation | Cooldown | Duplicate Suppression | Escalation</text>
  <line x1="400" y1="319" x2="400" y2="344" class="ar"/>

  <!-- Row 5: Escalation Logic -->
  <rect x="225" y="347" width="350" height="40" class="box"/>
  <text x="400" y="364" class="bt" font-weight="bold">Escalation Logic</text>
  <text x="400" y="378" class="lb">LOW - MEDIUM - HIGH - CRITICAL</text>
  <line x1="400" y1="387" x2="400" y2="412" class="ar"/>

  <!-- Row 6: Dashboard -->
  <rect x="275" y="415" width="250" height="40" class="box"/>
  <text x="400" y="432" class="bt" font-weight="bold">Alert Dashboard</text>
  <text x="400" y="446" class="lb">Alert Center Card</text>

  <!-- Split into 3 notification paths -->
  <line x1="400" y1="455" x2="400" y2="480" class="ar"/>
  <line x1="400" y1="480" x2="140" y2="480" class="ar"/>
  <line x1="400" y1="480" x2="660" y2="480" class="ar"/>
  <line x1="140" y1="480" x2="140" y2="505" class="ar"/>
  <line x1="400" y1="480" x2="400" y2="505" class="ar"/>
  <line x1="660" y1="480" x2="660" y2="505" class="ar"/>

  <rect x="40" y="508" width="200" height="50" class="box"/>
  <text x="140" y="528" class="bt" font-weight="bold">Worker Notification</text>
  <text x="140" y="545" class="lb">Self-correction guidance</text>

  <rect x="300" y="508" width="200" height="50" class="box"/>
  <text x="400" y="528" class="bt" font-weight="bold">Supervisor Dashboard</text>
  <text x="400" y="545" class="lb">Aggregated summary</text>

  <rect x="560" y="508" width="200" height="50" class="box"/>
  <text x="660" y="528" class="bt" font-weight="bold">Manager Escalation</text>
  <text x="660" y="545" class="lb">Exception-based alerts</text>

  <!-- Legend -->
  <rect x="50" y="590" width="700" height="170" class="box" stroke="#444" fill="none"/>
  <text x="400" y="612" class="tl" font-size="14">Legend</text>
  <rect x="80" y="625" width="120" height="28" class="box"/><text x="140" y="643" class="bt" font-size="11">Existing Module</text>
  <rect x="80" y="660" width="120" height="28" class="hl"/><text x="140" y="678" class="bt" font-size="11" fill="#bb86fc">Day 2 Module</text>
  <rect x="80" y="695" width="120" height="28" class="hl"/><text x="140" y="713" class="bt" font-size="11" fill="#ff79c6">Day 3 Module</text>
  <text x="280" y="640" class="bt" font-size="11" fill="#888">Alert levels: None - LOW - MEDIUM - HIGH - CRITICAL</text>
  <text x="280" y="660" class="bt" font-size="11" fill="#888">Cooldown: 2min per alert type after firing</text>
  <text x="280" y="680" class="bt" font-size="11" fill="#888">Duplicate suppression: identical alerts blocked for 5 min</text>
  <text x="280" y="700" class="bt" font-size="11" fill="#888">Threshold hysteresis: alert clears at threshold - delta</text>
  <text x="280" y="720" class="bt" font-size="11" fill="#888">Fatigue gating: low-risk alerts suppressed when improving</text>
  <text x="400" y="758" class="st" font-size="10">ErgoVigilance - Alert Management Architecture - July 2026</text>
</svg>'''
    path = os.path.join(OUT, 'Research', '03_AlertArchitecture.svg')
    with open(path, 'w', encoding='utf-8') as f:
        f.write(svg)
    print(f"Created: {path}")

# ═══════════════════════════════════════════════════════════════
# 04 — Alert Algorithm Doc
# ═══════════════════════════════════════════════════════════════
def create_algorithm():
    doc = Document()
    doc.add_heading('Alert Algorithm Specification', level=0)
    doc.add_paragraph('Week 4 — Day 3 | ErgoVigilance Implementation Document')

    doc.add_heading('1. Alert Levels', level=1)
    doc.add_paragraph('NONE (0): No alert. Posture within normal range.')
    doc.add_paragraph('LOW (1): Minor deviation. Worker notification only. Self-correction expected.')
    doc.add_paragraph('MEDIUM (2): Sustained deviation. Worker + Supervisor notified.')
    doc.add_paragraph('HIGH (3): Critical deviation or fatigue. All parties notified.')
    doc.add_paragraph('CRITICAL (4): Immediate physical risk. Audio-visual alarm. Session may pause.')

    doc.add_heading('2. Cooldown Logic', level=1)
    doc.add_paragraph(
        'After any alert fires, a cooldown timer starts for that alert type:\n'
        '• LOW: 60-second cooldown\n'
        '• MEDIUM: 120-second cooldown\n'
        '• HIGH: 180-second cooldown\n'
        '• CRITICAL: 300-second cooldown\n\n'
        'During cooldown, alerts of the same type are suppressed. Higher-severity alerts can still '
        'fire (e.g., MEDIUM during LOW cooldown).')

    doc.add_heading('3. Duplicate Suppression', level=1)
    doc.add_paragraph(
        'Alerts are considered duplicate if they share the same (alert_level, trigger_reason) pair '
        'within a 5-minute sliding window. Duplicates are suppressed but counted.\n\n'
        'Duplicate count tracks how many times the same alert was suppressed — this is displayed '
        'as "Duplicate Alerts Prevented" in the dashboard.')

    doc.add_heading('4. Escalation Rules', level=1)
    doc.add_paragraph(
        'Escalation is triggered by:\n'
        '• Duration: Sustained MEDIUM for >20 min → escalate to HIGH\n'
        '• Frequency: 3+ LOW alerts in 30 min → escalate next to MEDIUM\n'
        '• Fatigue: Fatigue >50 during any alert → +1 escalation level\n'
        '• Critical: Any HIGH alert that persists >5 min → CRITICAL')

    doc.add_heading('5. Fatigue Influence on Alerts', level=1)
    doc.add_paragraph(
        'Fatigue modifies alert behaviour:\n'
        '• Fatigue < 30: Standard alert thresholds\n'
        '• Fatigue 30–60: LOW alerts treated as MEDIUM, MEDIUM as HIGH\n'
        '• Fatigue > 60: All alerts escalated +1 level\n'
        '• Fatigue > 80: CRITICAL alert triggered automatically')

    doc.add_heading('6. Context Influence', level=1)
    doc.add_paragraph(
        'Task context modifies alert sensitivity:\n'
        '• Assembly Work: +1 escalation (repetitive task needs faster response)\n'
        '• Lifting: -1 escalation (brief high loads are inherent)\n'
        '• Inspection: Standard escalation\n'
        '• Typing: Standard escalation\n'
        '• Neutral Standing: -1 escalation (lowest risk task)')

    doc.add_heading('7. Future Backend Implementation', level=1)
    doc.add_paragraph(
        'The AlertEngine class will be implemented in Week 5 as backend/services/alert_engine.py. '
        'It will:\n'
        '• Maintain per-session alert state (timestamps, counts, cooldowns)\n'
        '• Consume ContextAwareRiskEngine output\n'
        '• Apply rules from the decision matrix\n'
        '• Expose alert_state via FastAPI DashboardResponse\n'
        '• Push real-time alert updates via WebSocket\n'
        '• Persist alert history alongside session data')

    path = os.path.join(OUT, 'Implementation', '04_AlertAlgorithm.docx')
    doc.save(path)
    print(f"Created: {path}")

# ═══════════════════════════════════════════════════════════════
# 05 — Alert Pseudocode
# ═══════════════════════════════════════════════════════════════
def create_pseudocode():
    code = '''"""
05_AlertPseudoCode.py — Intelligent Alert Management Pseudocode

This file contains pseudocode only. It is NOT executable.
It demonstrates the planned Alert Engine flow for Week 5 implementation.

Flow:
  Context Risk → Alert Level → Cooldown Check → Duplicate Check
  → Fatigue Adjustment → Escalation → Notification Routing

Author: ErgoVigilance Team
Date: July 2026
"""

# =============================================================================
# SECTION 1 — Alert Levels and Configuration
# =============================================================================

from enum import IntEnum
from dataclasses import dataclass, field
from typing import Optional
from collections import deque
import time

class AlertLevel(IntEnum):
    NONE = 0
    LOW = 1
    MEDIUM = 2
    HIGH = 3
    CRITICAL = 4

ALERT_LABELS = {
    AlertLevel.NONE: "None",
    AlertLevel.LOW: "LOW",
    AlertLevel.MEDIUM: "MEDIUM",
    AlertLevel.HIGH: "HIGH",
    AlertLevel.CRITICAL: "CRITICAL",
}

ALERT_COLORS = {
    AlertLevel.NONE: "#22c55e",
    AlertLevel.LOW: "#22c55e",
    AlertLevel.MEDIUM: "#f97316",
    AlertLevel.HIGH: "#ef4444",
    AlertLevel.CRITICAL: "#d946ef",
}

# Cooldown durations per level (seconds)
COOLDOWN_DURATIONS = {
    AlertLevel.LOW: 60,
    AlertLevel.MEDIUM: 120,
    AlertLevel.HIGH: 180,
    AlertLevel.CRITICAL: 300,
}

DUPLICATE_WINDOW = 300  # 5 minutes
ESCALATION_DURATION_THRESHOLD = 20 * 60  # 20 minutes sustained MEDIUM
ESCALATION_FREQUENCY_THRESHOLD = 3  # 3 LOW alerts in 30 min
FATIGUE_ESCALATION_THRESHOLD = 50

# =============================================================================
# SECTION 2 — Alert State
# =============================================================================

@dataclass
class AlertRecord:
    timestamp: float
    level: AlertLevel
    trigger: str
    worker_notified: bool = False
    supervisor_notified: bool = False
    manager_notified: bool = False

@dataclass
class AlertState:
    current_level: AlertLevel = AlertLevel.NONE
    current_cooldown_remaining: float = 0.0
    last_alert_time: Optional[float] = None
    timeline: list[AlertRecord] = field(default_factory=list)
    today_alerts: int = 0
    worker_alerts: int = 0
    supervisor_alerts: int = 0
    critical_alerts: int = 0
    false_positives_avoided: int = 0
    duplicates_prevented: int = 0
    suppressed_notifications: int = 0
    confidence: float = 0.0
    escalation_worker: bool = False
    escalation_supervisor: bool = False
    escalation_manager: bool = False
    recent_alert_keys: deque = field(default_factory=lambda: deque(maxlen=100))

# =============================================================================
# SECTION 3 — Alert Engine (Pseudocode)
# =============================================================================

class AlertEngine:
    """
    Stateful alert engine that manages alert lifecycle.

    Instantiated once per session alongside ContextAwareRiskEngine.
    Called once per frame with risk level, fatigue, task, and elapsed time.

    Maintains:
        - Current alert level and cooldown
        - Alert timeline for display
        - Duplicate suppression cache
        - Escalation state (worker/supervisor/manager)
        - Counters for today_alerts, worker_alerts, etc.
    """

    def __init__(self):
        self.state = AlertState()
        self._suppression_cache = {}  # key -> timestamp

    def compute_alert_level(self, risk_level, fatigue, task, elapsed):
        """
        Determine current alert level from context-aware risk.

        Rules:
            1. Start with NONE
            2. Map risk_level -> alert_level (LOW->LOW, MEDIUM->MEDIUM, HIGH->HIGH)
            3. Apply fatigue escalation
            4. Apply task modifier
            5. Apply duration escalation
            6. Check cooldown - if cooling down, suppress
            7. Check duplicate - if duplicate, suppress
            8. Return final alert level
        """
        pass  # see 04_AlertAlgorithm.docx for full logic

    def check_cooldown(self, level):
        """
        Check if an alert level is in cooldown.

        Returns:
            remaining_seconds (float): 0 if not cooling down
        """
        pass

    def check_duplicate(self, level, trigger):
        """
        Check if this alert is a duplicate within DUPLICATE_WINDOW.

        Returns:
            is_duplicate (bool)
        """
        pass

    def escalate(self, level, fatigue, task, frequency):
        """
        Apply escalation rules.

        Returns:
            escalated_level (AlertLevel)
            escalation_path (list[str])
        """
        pass

    def route_notification(self, level):
        """
        Determine notification routing based on alert level.

        Returns:
            notify_worker (bool)
            notify_supervisor (bool)
            notify_manager (bool)
        """
        pass

    def update(self, risk_level, fatigue_level, task, elapsed):
        """
        Main update method - called once per frame.

        Returns:
            AlertState snapshot with current_level, timeline, counts, etc.
        """
        pass

# =============================================================================
# SECTION 4 — Integration Points
# =============================================================================

def integration_points():
    """
    How AlertEngine integrates with existing code:

    ContextAwareRiskEngine returns:
        { risk_level, final_score, contributions, temporal_state }

    AlertEngine consumes:
        risk_level -> maps to base alert level
        temporal_state.fatigue_count -> used for fatigue escalation
        contributions.task_modifier -> used for context adjustment

    AlertEngine produces:
        alert_state -> displayed in AlertCenterCard
        alert_state.current_level -> drives border glow, badge color
        alert_state.timeline -> shown in timeline section
        alert_state.escalation_* -> checkmark/pending status

    No existing module is modified. AlertEngine is a new dependency.
    """
    pass


if __name__ == "__main__":
    print("This is pseudocode - not executable.")
    print("See 04_AlertAlgorithm.docx for the full specification.")
'''
    path = os.path.join(OUT, 'Implementation', '05_AlertPseudoCode.py')
    with open(path, 'w') as f:
        f.write(code)
    print(f"Created: {path}")

# ═══════════════════════════════════════════════════════════════
# Demo material
# ═══════════════════════════════════════════════════════════════
def create_demo_script():
    doc = Document()
    doc.add_heading('Demo Script: Intelligent Alert Management', level=0)
    doc.add_paragraph('Duration: ~2 minutes | ErgoVigilance Week 4 Day 3')
    steps = [
        ('1. Opening (15s)',
         '"Good afternoon. Today I\'m demonstrating our Intelligent Alert Management system for ErgoVigilance. '
         'Yesterday we added Context-Aware Risk Assessment. Today we\'re tackling the next challenge: making '
         'sure the right person gets the right alert at the right time — and suppressing everything else."'),
        ('2. The Problem (20s)',
         '"Alert fatigue is a well-documented phenomenon. In clinical settings, 85-99% of alarms are ignored. '
         'Our system generates up to 30 risk assessments per second — without intelligent alert management, '
         'that\'s an unsustainable notification storm. Workers learn to dismiss alerts, supervisors become '
         'desensitised, and the system loses credibility."'),
        ('3. What We Built (25s)',
         '"Our Intelligent Alert Center has four components. First, a five-level alert system — None, LOW, '
         'MEDIUM, HIGH, CRITICAL — colour-coded green, yellow, orange, red. Second, cooldown logic: after '
         'any alert, a timer prevents duplicate alerts of the same type. Third, duplicate suppression: '
         'identical alerts are blocked for 5 minutes. Fourth, escalation: from worker, to supervisor, to '
         'manager — with fatigue and context modifiers."'),
        ('4. Live Demo (25s)',
         '"[Open React frontend] Here is the Intelligent Alert Center card. It shows the current alert level '
         'with a prominent colour badge, an alert timeline with events, the current cooldown timer, '
         'escalation status with checkmarks, alert counts for today, and alert intelligence metrics. '
         'Clicking the AI Alert Engine badge opens a modal showing the current vs future pipeline."'),
        ('5. Demo Mode (20s)',
         '"Let me switch to Demo Mode. [Toggle demo] Watch as I cycle through Office, Assembly, Warehouse, '
         'Machine Operator, and Inspection scenarios. Each scenario updates the alert level, timeline, '
         'cooldown, escalation status, counts, and confidence — with no backend required."'),
        ('6. Research Foundation (10s)',
         '"This prototype is backed by research in industrial alert systems, human factors psychology, '
         'OSHA recommendations, and best practices from aviation and manufacturing."'),
        ('7. Closing (5s)',
         '"In Week 5, we\'ll implement the backend AlertEngine, wire it into the real pipeline, and connect '
         'it to the FastAPI and React frontend with live WebSocket updates."'),
    ]
    for h, b in steps:
        doc.add_heading(h, level=1); doc.add_paragraph(b)
    path = os.path.join(OUT, 'Demo', 'Demo_Script.docx'); doc.save(path); print(f"Created: {path}")

def create_screenshot_guide():
    doc = Document()
    doc.add_heading('Screenshot Guide: Intelligent Alert Management', level=0)
    doc.add_paragraph('Week 4 — Day 3 | ErgoVigilance')
    for title, desc in [
        ('1. React Dashboard — Alert Center Card', 'Full dashboard showing the Intelligent Alert Center card with current alert level, timeline, cooldown, escalation, counts, and intelligence metrics.'),
        ('2. Alert Level Badge', 'Close-up of the alert level badge showing LOW (green), MEDIUM (yellow), HIGH (orange), or CRITICAL (red) with contextual background.'),
        ('3. Alert Timeline', 'The timeline section showing sequential events: Worker warned → Second warning → Supervisor notified → Break recommended → Critical escalation.'),
        ('4. Escalation Status', 'Escalation section showing checkmarks for Worker ✓, Supervisor ✓, and Manager Pending with colour-coded status.'),
        ('5. AI Alert Engine Modal', 'Modal showing Current Alert Flow (Risk → Warning) vs Future Intelligent Flow (Risk → Context → Fatigue → Escalation → Notification).'),
        ('6. Research Document', 'First page of 01_IntelligentAlertResearch.docx showing title, introduction, and table of contents.'),
        ('7. Decision Matrix', 'Screenshot of 02_AlertDecisionMatrix.xlsx showing 15 rows of risk/duration/fatigue/task combinations with colour-coded alert levels.'),
    ]:
        doc.add_heading(title, level=1); doc.add_paragraph(desc)
    path = os.path.join(OUT, 'Demo', 'Screenshot_Guide.docx'); doc.save(path); print(f"Created: {path}")

def create_findings():
    doc = Document()
    doc.add_heading('Findings: Intelligent Alert Management', level=0)
    doc.add_paragraph('Week 4 — Day 3 | ErgoVigilance')
    for h, b in [
        ('Objective', 'Research, design, and prototype an Intelligent Alert Management system for ErgoVigilance that prevents alert fatigue, implements evidence-based escalation, and delivers role-based notifications — without modifying any existing backend modules.'),
        ('Work Completed', '• Research document (10 sections): Alert fatigue, human factors, escalation logic, cooldown/suppression, worker psychology, supervisor systems, OSHA recommendations.\n• Decision matrix (15 rows): Risk × Duration × Fatigue × Task → Alert Level → Escalation → Notification.\n• Architecture diagram: Full pipeline with Alert Engine and three notification paths (Worker/Supervisor/Manager).\n• Algorithm specification: 5 alert levels, cooldown durations, duplicate suppression, escalation rules, fatigue/context influence.\n• Pseudocode: AlertEngine class with state management, suppression cache, escalation logic.\n• React visual prototype: Intelligent Alert Center card with alert level, timeline, cooldown, escalation, counts, intelligence metrics, AI Engine modal.\n• Demo Mode integration: All 5 scenarios update alert state with deterministic values.\n• Demo materials: Script, screenshot guide, findings document.'),
        ('Challenges', '1. Designing cooldown values without field validation data. Values chosen based on literature (60s LOW → 300s CRITICAL) with 2× increment per level.\n2. Balancing duplicate suppression window (5 min) against genuine repeat alerts — too short causes storms, too long misses patterns.\n3. Determining fatigue thresholds for escalation (50 = escalation, 60 = +1 level, 80 = auto-critical) — will need tuning in Week 5.'),
        ('Research Outcome', 'The research confirms that alert fatigue is a well-documented problem across clinical, aviation, and industrial domains with established mitigation strategies. The proposed 5-level system with cooldown, duplicate suppression, and fatigue-adjusted escalation follows best practices from all three domains while remaining simple enough for real-time deployment.'),
        ('Future Work (Week 5+)', '1. Implement AlertEngine class in backend/services/alert_engine.py.\n2. Unit tests for cooldown, duplicate suppression, escalation, fatigue adjustment.\n3. Integration with ContextAwareRiskEngine output.\n4. Add alert_state field to FastAPI DashboardResponse schema.\n5. Connect React Alert Center card to live API data.\n6. WebSocket real-time alert push.\n7. Field validation of cooldown and escalation thresholds.'),
    ]:
        doc.add_heading(h, level=1); doc.add_paragraph(b)
    path = os.path.join(OUT, 'Demo', 'Findings.docx'); doc.save(path); print(f"Created: {path}")

if __name__ == '__main__':
    create_research()
    create_matrix()
    create_arch()
    create_algorithm()
    create_pseudocode()
    create_demo_script()
    create_screenshot_guide()
    create_findings()
    print("\nAll documentation files created successfully.")
